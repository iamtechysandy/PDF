import streamlit as st
from difflib import unified_diff
from fuzzywuzzy import fuzz
from fuzzywuzzy import process
from docx import Document
from PyPDF2 import PdfReader
import io
import matplotlib.pyplot as plt


def run():
    st.title("Enhanced Document Comparison Tool")
    st.markdown("### Upload two documents (PDF, Word, or Text) to compare their content and formatting.")
    st.markdown("---")

    # File uploaders
    col1, col2 = st.columns(2)
    with col1:
        file1 = st.file_uploader("Upload First Document", type=["pdf", "docx", "txt"])
    with col2:
        file2 = st.file_uploader("Upload Second Document", type=["pdf", "docx", "txt"])

    if file1 and file2:
        # Extract text from the files
        file1_type = file1.type.split("/")[1]
        file2_type = file2.type.split("/")[1]

        text1 = extract_text(file1, file1_type)
        text2 = extract_text(file2, file2_type)

        if not text1 or not text2:
            st.error("Could not extract text from one or both files. Please ensure the files are valid.")
        else:
            # Compare texts with fuzzy matching
            st.subheader("Content Comparison")
            matched, unmatched = fuzzy_compare_lines(text1, text2)

            st.write(f"**Matched Lines ({len(matched)})**")
            for m in matched[:10]:
                st.write(f"File1: {m[0]} | File2: {m[1]} | Similarity: {m[2]}%")

            st.write(f"**Unmatched Lines ({len(unmatched)})**")
            for u in unmatched[:10]:
                st.write(f"File1: {u[0]} | Similarity Score: {u[2]}%")

            # Word formatting comparison
            if file1_type == "docx" and file2_type == "docx":
                st.subheader("Formatting Comparison")
                formatting_differences = compare_word_formatting(file1, file2)
                if formatting_differences:
                    for diff in formatting_differences:
                        st.write(diff)
                else:
                    st.write("No formatting differences found.")

            # Visualization of comparison
            st.subheader("Comparison Summary")
            total_lines = len(text1.splitlines()) + len(text2.splitlines())
            matched_count = len(matched)
            unmatched_count = len(unmatched)

            # Pie chart
            fig, ax = plt.subplots()
            ax.pie(
                [matched_count, unmatched_count],
                labels=["Matched Lines", "Unmatched Lines"],
                autopct="%1.1f%%",
                colors=["#4CAF50", "#FF5733"],
            )
            st.pyplot(fig)

            # Option to download the differences
            diff_output = "\n".join([f"Matched: {m}" for m in matched] + [f"Unmatched: {u}" for u in unmatched])
            output_buffer = io.BytesIO()
            output_buffer.write(diff_output.encode("utf-8"))
            output_buffer.seek(0)

            st.download_button(
                "Download Comparison Report",
                data=output_buffer,
                file_name="comparison_report.txt",
                mime="text/plain",
            )


# Helper Functions

def extract_text(file, file_type):
    """
    Extract text from PDF, Word, or text files.
    """
    if file_type == "pdf":
        reader = PdfReader(file)
        text = "\n".join(page.extract_text() for page in reader.pages if page.extract_text())
    elif file_type == "docx":
        doc = Document(file)
        text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    elif file_type == "txt":
        text = file.getvalue().decode("utf-8")
    else:
        text = ""
    return text


def fuzzy_compare_lines(text1, text2, threshold=80):
    """
    Perform fuzzy matching line by line.
    """
    lines1 = text1.splitlines()
    lines2 = text2.splitlines()
    matched = []
    unmatched = []

    for line in lines1:
        match, score = process.extractOne(line, lines2)
        if score >= threshold:
            matched.append((line, match, score))
        else:
            unmatched.append((line, None, score))
    return matched, unmatched


def compare_word_formatting(file1, file2):
    """
    Compare formatting styles in Word documents.
    """
    doc1 = Document(file1)
    doc2 = Document(file2)

    formats = []
    for para1, para2 in zip(doc1.paragraphs, doc2.paragraphs):
        if para1.style.name != para2.style.name:
            formats.append(f"Mismatch in styles: '{para1.text}' vs '{para2.text}'")
    return formats
